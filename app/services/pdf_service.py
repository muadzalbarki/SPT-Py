import subprocess
import sys
import os
import tempfile
import shutil
from pathlib import Path
from typing import Optional
from PySide6.QtCore import QObject, Signal

from app.settings import AppSettings
from app.services.log_service import LogService


class PdfService(QObject):
    conversion_finished = Signal(str)
    conversion_error = Signal(str)

    ENGINE_AUTO = "auto"
    ENGINE_MSWORD = "msword"
    ENGINE_LIBREOFFICE = "libreoffice"

    _soffice_name = "soffice.exe" if sys.platform == "win32" else "soffice"
    _word_available: Optional[bool] = None
    _libreoffice_available: Optional[bool] = None

    def __init__(self, parent=None):
        super().__init__(parent)
        self._logger = LogService.get_logger("pdf_service")

    def convert_to_pdf(self, docx_path: str, output_dir: Optional[str] = None) -> Optional[str]:
        docx_path = str(Path(docx_path).resolve())
        if output_dir is None:
            output_dir = str(Path(docx_path).parent)

        engine = self._get_engine()
        self._logger.info("PDF conversion engine: %s", engine)

        if engine in (self.ENGINE_AUTO, self.ENGINE_MSWORD):
            result = self._convert_with_word(docx_path, output_dir)
            if result is not None:
                return result
            if engine == self.ENGINE_MSWORD:
                msg = "Microsoft Word gagal mengkonversi dokumen"
                self._logger.error(msg)
                self.conversion_error.emit(msg)
                return None

        return self._convert_with_libreoffice(docx_path, output_dir)

    def _convert_with_word(self, docx_path: str, output_dir: str) -> Optional[str]:
        if not self.is_word_available():
            return None
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.DisplayAlerts = False

            doc = word.Documents.Open(docx_path)
            pdf_path = str(Path(output_dir) / (Path(docx_path).stem + ".pdf"))
            doc.SaveAs(pdf_path, FileFormat=17)
            doc.Close()
            word.Quit()

            if Path(pdf_path).exists():
                self._logger.info("Word PDF: %s", pdf_path)
                self.conversion_finished.emit(pdf_path)
                return pdf_path
        except Exception as e:
            self._logger.error("Word conversion error: %s", e)
        return None

    def _convert_with_libreoffice(self, docx_path: str, output_dir: str) -> Optional[str]:
        output_pdf = Path(output_dir) / (Path(docx_path).stem + ".pdf")
        try:
            result = subprocess.run(
                self._get_soffice_cmd("--headless", "--convert-to", "pdf", "--outdir", output_dir, docx_path),
                capture_output=True,
                text=True,
                timeout=60,
            )
            if result.returncode == 0 and output_pdf.exists():
                self._logger.info("LO PDF: %s", output_pdf)
                self.conversion_finished.emit(str(output_pdf))
                return str(output_pdf)
            error_msg = result.stderr or "LibreOffice conversion failed"
            self._logger.error("LO error: %s", error_msg)
            self.conversion_error.emit(error_msg)
            return None
        except FileNotFoundError:
            msg = "LibreOffice tidak ditemukan"
            self._logger.error(msg)
            self.conversion_error.emit(msg)
            return None
        except subprocess.TimeoutExpired:
            msg = "Konversi PDF timeout (60 detik)"
            self._logger.error(msg)
            self.conversion_error.emit(msg)
            return None
        except Exception as e:
            self._logger.error("LO error: %s", e)
            self.conversion_error.emit(str(e))
            return None

    def test_export(self) -> tuple:
        tmp_dir = Path(tempfile.mkdtemp())
        tmp_docx = tmp_dir / "test_export.docx"
        try:
            from docx import Document
            doc = Document()
            doc.add_paragraph("Test PDF Export - SPT-DPRD")
            doc.add_paragraph("Dokumen ini dibuat untuk verifikasi engine PDF.")
            doc.save(str(tmp_docx))
            result = self.convert_to_pdf(str(tmp_docx), str(tmp_dir))
            if result:
                return True, "Export berhasil"
            return False, "Export gagal — lihat log untuk detail"
        except Exception as e:
            return False, f"Error: {e}"
        finally:
            shutil.rmtree(tmp_dir, ignore_errors=True)

    def is_available(self) -> bool:
        return self.is_word_available() or self.is_libreoffice_available()

    @classmethod
    def is_word_available(cls) -> bool:
        if cls._word_available is not None:
            return cls._word_available
        if sys.platform != "win32":
            cls._word_available = False
            return False
        try:
            import win32com.client
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            word.Quit()
            cls._word_available = True
        except Exception:
            cls._word_available = False
        return cls._word_available

    @classmethod
    def is_libreoffice_available(cls) -> bool:
        if cls._libreoffice_available is not None:
            return cls._libreoffice_available
        try:
            result = subprocess.run(
                cls._get_soffice_cmd("--version"),
                capture_output=True, text=True, timeout=10,
            )
            cls._libreoffice_available = result.returncode == 0
        except (FileNotFoundError, subprocess.TimeoutExpired):
            cls._libreoffice_available = False
        return cls._libreoffice_available

    @classmethod
    def _get_soffice_cmd(cls, *args: str) -> list[str]:
        return [cls._soffice_name, *args]

    def _get_engine(self) -> str:
        return AppSettings.instance().get("pdf_engine", self.ENGINE_AUTO)

    @classmethod
    def reset_cache(cls):
        cls._word_available = None
        cls._libreoffice_available = None
