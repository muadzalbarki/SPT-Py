import re
from pathlib import Path
from typing import Any
from docx import Document


class DraftParserService:
    @staticmethod
    def parse_employee_draft(docx_path: str) -> list[dict[str, Any]]:
        path = Path(docx_path)
        if not path.exists():
            return []

        document = Document(str(path))
        entries: list[dict[str, Any]] = []
        lines: list[str] = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if not text:
                continue
            if re.search(r"\s[-–—]\s", text):
                parts = re.split(r"\s[-–—]\s", text, maxsplit=1)
                if len(parts) == 2:
                    entries.append({
                        "nama": parts[0].strip(),
                        "jabatan": parts[1].strip(),
                        "komisi": "",
                    })
                    continue
            lines.append(text)

        for table in document.tables:
            if not table.rows:
                continue
            headers = [cell.text.strip().lower() for cell in table.rows[0].cells]
            if any("nama" in header for header in headers) and any("jabatan" in header for header in headers):
                nama_index = next(i for i, header in enumerate(headers) if "nama" in header)
                jabatan_index = next(i for i, header in enumerate(headers) if "jabatan" in header)
                komisi_index = next((i for i, header in enumerate(headers) if "komisi" in header), None)
                for row in table.rows[1:]:
                    nama = row.cells[nama_index].text.strip()
                    jabatan = row.cells[jabatan_index].text.strip()
                    komisi = row.cells[komisi_index].text.strip() if komisi_index is not None else ""
                    if nama:
                        entries.append({
                            "nama": nama,
                            "jabatan": jabatan or "Anggota",
                            "komisi": komisi,
                        })

        current: dict[str, Any] = {}
        for line in lines:
            nama_match = re.match(r"(?i)nama\s*[:\-]\s*(.+)", line)
            jabatan_match = re.match(r"(?i)jabatan\s*[:\-]\s*(.+)", line)
            komisi_match = re.match(r"(?i)komisi\s*[:\-]\s*(.+)", line)
            if nama_match:
                if current.get("nama"):
                    entries.append(current)
                    current = {}
                current["nama"] = nama_match.group(1).strip()
                continue
            if jabatan_match and current:
                current["jabatan"] = jabatan_match.group(1).strip()
                continue
            if komisi_match and current:
                current["komisi"] = komisi_match.group(1).strip()
                continue
            if re.match(r"^(.+?),\s*(.+)$", line):
                parts = [part.strip() for part in line.split(",", 1)]
                if len(parts) == 2 and parts[0] and parts[1]:
                    entries.append({"nama": parts[0], "jabatan": parts[1], "komisi": ""})

        if current.get("nama"):
            current.setdefault("jabatan", "Anggota")
            current.setdefault("komisi", "")
            entries.append(current)

        unique = []
        seen = set()
        for item in entries:
            key = (item.get("nama", "").lower(), item.get("jabatan", "").lower())
            if key not in seen:
                seen.add(key)
                unique.append(item)

        return unique
