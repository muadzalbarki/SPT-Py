import re
import os
import subprocess
from pathlib import Path
from typing import Any
from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm
from docx import Document


PLACEHOLDER_PATTERN = re.compile(r"\{\s*([a-zA-Z0-9_\.]+)\s*\}")


class TemplateEngine:
    @staticmethod
    def extract_placeholders(template_path: str) -> set[str]:
        path = Path(template_path)
        if not path.exists():
            return set()

        document = Document(str(path))
        placeholders = set()

        def scan_text(text: str) -> None:
            for match in PLACEHOLDER_PATTERN.findall(text):
                placeholders.add(match)

        for paragraph in document.paragraphs:
            scan_text(paragraph.text)
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    scan_text(cell.text)

        return placeholders

    @staticmethod
    def render_template(template_path: str, context: dict[str, Any], output_path: str) -> str:
        template = DocxTemplate(template_path)
        template.render(context)
        output_dir = Path(output_path).parent
        output_dir.mkdir(parents=True, exist_ok=True)
        template.save(output_path)
        return output_path

    @staticmethod
    def convert_to_pdf(docx_path: str, output_dir: str) -> str:
        output_dir_path = Path(output_dir)
        output_dir_path.mkdir(parents=True, exist_ok=True)
        command = [
            "soffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(output_dir_path),
            str(docx_path),
        ]
        try:
            subprocess.run(command, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        except FileNotFoundError:
            command[0] = "libreoffice"
            subprocess.run(command, check=True, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)

        target_pdf = output_dir_path / f"{Path(docx_path).stem}.pdf"
        return str(target_pdf)

    @staticmethod
    def write_sample_template(path: str) -> None:
        if Path(path).exists():
            return

        doc = Document()
        doc.add_paragraph("PEMERINTAH DAERAH DPRD")
        doc.add_paragraph("SURAT TUGAS")
        doc.add_paragraph("Nomor: {nomor_surat}")
        doc.add_paragraph("Tanggal: {tanggal}")
        doc.add_paragraph("Kepada Yth:")
        doc.add_paragraph("{nama}")
        doc.add_paragraph("Jabatan: {jabatan}")
        doc.add_paragraph("NIP: {nip}")
        doc.add_paragraph("Komisi: {komisi}")
        doc.add_paragraph("\nDaftar Peserta:")
        doc.add_paragraph("{% for peserta in peserta %}")
        doc.add_paragraph("- {{ peserta.nama }} / {{ peserta.jabatan }}")
        doc.add_paragraph("{% endfor %}")
        doc.save(path)
